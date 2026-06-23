"""
upload_routes.py
----------------
API endpoint for file uploading:
- POST /upload  → accepts PDF, DOCX, or TXT files
                  extracts their text
                  ingests text into the RAG vector store
                  so the chatbot can answer questions about the file

Supported formats: PDF (.pdf), Word (.docx), Plain text (.txt)
"""

import io
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from backend.auth import decode_token
from backend.rag.vectorstore import save_vectorstore
from langchain_core.documents import Document

upload_router = APIRouter(prefix="/upload")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Pull plain text out of PDF, DOCX, or TXT files."""
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)

    elif ext == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")

    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT.")


@upload_router.post("")
async def upload_file(file: UploadFile = File(...), user_id: str = Depends(decode_token)):
    file_bytes = await file.read()
    text = extract_text(file_bytes, file.filename)

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the file.")

    # Wrap text in a LangChain Document and push into the RAG vector store
    # Split into ~1000-char chunks so embeddings stay manageable
    chunk_size = 1000
    chunks = [
        Document(
            page_content=text[i:i + chunk_size],
            metadata={"source": file.filename, "uploaded_by": user_id}
        )
        for i in range(0, len(text), chunk_size)
    ]

    save_vectorstore(chunks)

    return {
        "message": f"'{file.filename}' uploaded and indexed successfully.",
        "chunks_indexed": len(chunks)
    }
