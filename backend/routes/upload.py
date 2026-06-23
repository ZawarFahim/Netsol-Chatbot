import io
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from backend.services.auth import decode_token
from backend.rag.vectorstore import save_vectorstore
from langchain_core.documents import Document

upload_router = APIRouter(prefix="/upload")

def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        import fitz
        return "\n".join(page.get_text() for page in fitz.open(stream=file_bytes, filetype="pdf"))
    elif ext == "docx":
        from docx import Document as DocxDocument
        return "\n".join(p.text for p in DocxDocument(io.BytesIO(file_bytes)).paragraphs if p.text.strip())
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    raise HTTPException(status_code=400, detail=f"Unsupported file format: .{ext}. Please upload a PDF, DOCX, or TXT file.")

@upload_router.post("")
async def upload_file(file: UploadFile = File(...), user_id: str = Depends(decode_token)):
    text = extract_text(await file.read(), file.filename)
    if not text.strip():
        raise HTTPException(status_code=400, detail="The document appears to be empty or contains no readable text (it might be a scanned image).")

    chunks = [
        Document(page_content=text[i:i + 1000], metadata={"source": file.filename, "uploaded_by": user_id})
        for i in range(0, len(text), 1000)
    ]
    save_vectorstore(chunks)
    return {"message": "Success"}
